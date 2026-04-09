from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
import pdfplumber
import docx
import io
import json
from typing import Optional
from ai_service import analyze_resume_genuine, improve_resume_ai, parse_resume_ai

app = FastAPI(title="ATS Resume Checker AI-Engine", version="2.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    text = ""
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Could not read PDF: {str(e)}")
    return text.strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    text = ""
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Could not read DOCX: {str(e)}")
    return text.strip()


@app.post("/api/upload-resume")
async def upload_resume(file: UploadFile = File(...)):
    content = await file.read()
    if len(content) > 5 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File too large. Max 5MB allowed.")
    
    filename = file.filename.lower()
    if filename.endswith(".pdf"):
        text = extract_text_from_pdf(content)
    elif filename.endswith(".docx"):
        text = extract_text_from_docx(content)
    elif filename.endswith(".txt"):
        text = content.decode("utf-8", errors="ignore")
    else:
        raise HTTPException(status_code=400, detail="Unsupported file type. Use PDF, DOCX, or TXT.")
    
    if len(text.strip()) < 50:
        raise HTTPException(status_code=400, detail="Could not extract text. Try pasting text manually.")

    # Pure AI Parsing
    try:
        ai_info = parse_resume_ai(text)
    except Exception as e:
        print(f"AI Parsing Failed: {e}")
        ai_info = { "name": "User", "suggested_role": "", "summary": "", "sections_found": {} }

    return {
        "success": True,
        "filename": file.filename,
        "text": text,
        "word_count": len(text.split()),
        "char_count": len(text),
        "sections_detected": ai_info.get("sections_found", {}),
        "ai_info": ai_info
    }


@app.post("/api/analyze")
async def analyze_resume(
    resume_text: str = Form(...),
    job_role: str = Form(...),
    job_description: Optional[str] = Form(None)
):
    if len(resume_text.strip()) < 50:
        raise HTTPException(status_code=400, detail="Resume text too short.")
    if not job_role.strip():
        raise HTTPException(status_code=400, detail="Job role is required.")

    try:
        return analyze_resume_genuine(resume_text, job_role, job_description)
    except Exception as e:
        print(f"AI Analysis Failed: {e}")
        raise HTTPException(
            status_code=503,
            detail=f"AI analysis failed: {str(e)}. Please try again."
        )


@app.post("/api/improve")
async def improve_resume(
    resume_text: str = Form(...),
    job_role: str = Form(...),
    missing_keywords: str = Form(...),
    suggestions: str = Form(...),
    confirmed_skills: Optional[str] = Form(None)
):
    confirmed = json.loads(confirmed_skills) if confirmed_skills else []

    try:
        improved = improve_resume_ai(resume_text, job_role, confirmed)
        return {"success": True, "improved_text": improved}
    except Exception as e:
        print(f"AI Improve Failed: {e}")
        raise HTTPException(
            status_code=503,
            detail=f"AI resume improvement failed: {str(e)}."
        )


@app.get("/api/health")
def health():
    return {"status": "ok", "version": "2.0.0", "engine": "Gemini-Pure-AI"}


if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)